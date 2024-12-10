#Работу выполнил студент группы 44ИС-21 Казазян Эдгар
#Вариант 2. Автосалон.
#Менеджеры автосалона осуществляют продажу клиентам автомобилей различных марок.
#Таблицы: Покупатели (Код покупателя, ФИО), Менеджеры (Код менеджера, ФИО), Автомобили (код, марка), Продажи (Код менеджера, код автомобиля, государственный номер, код покупателя, дата, цена).
#Определить:
#- среднюю сумму сделки
#- долю продаж автомобилей разных марок.

import sys
import mysql.connector
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QTableWidget, QTableWidgetItem,
    QPushButton, QVBoxLayout, QWidget, QLabel, QLineEdit, QFormLayout, QComboBox, QMessageBox, QGridLayout
)
from docx import Document
import os
import matplotlib.pyplot as plt
from PyQt6.QtGui import QPixmap


class AutoSalonApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Автосалон")
        self.setGeometry(100, 100, 1000, 700)
        self.setStyleSheet("font-family: Arial; font-size: 14px; background-color: #e1f5d7;")
        self.connection = self.connect_to_db()
        self.init_ui()

    def connect_to_db(self):
        try:
            return mysql.connector.connect(
                host="localhost",
                user="root",
                password="",
                database="var2"
            )
        except mysql.connector.Error as e:
            QMessageBox.critical(self, "Ошибка подключения", f"Не удалось подключиться к базе данных: {e}")
            sys.exit()

    def init_ui(self):
        layout = QVBoxLayout()

        self.table = QTableWidget()
        layout.addWidget(self.table)

        buttons_layout = QGridLayout()

        self.add_customer_button = QPushButton("Добавить клиента")
        self.add_customer_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.add_customer_button.clicked.connect(self.add_customer)
        buttons_layout.addWidget(self.add_customer_button, 0, 0)

        self.add_manager_button = QPushButton("Добавить менеджера")
        self.add_manager_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.add_manager_button.clicked.connect(self.add_manager)
        buttons_layout.addWidget(self.add_manager_button, 0, 1)

        self.add_car_button = QPushButton("Добавить автомобиль")
        self.add_car_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.add_car_button.clicked.connect(self.add_car)
        buttons_layout.addWidget(self.add_car_button, 1, 0)

        self.add_sale_button = QPushButton("Добавить продажу")
        self.add_sale_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.add_sale_button.clicked.connect(self.add_sale)
        buttons_layout.addWidget(self.add_sale_button, 1, 1)

        self.avg_button = QPushButton("Средняя сумма сделки")
        self.avg_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.avg_button.clicked.connect(self.calculate_average)
        buttons_layout.addWidget(self.avg_button, 2, 0)

        self.share_button = QPushButton("Доля продаж по маркам")
        self.share_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.share_button.clicked.connect(self.calculate_share)
        buttons_layout.addWidget(self.share_button, 2, 1)

        self.export_button = QPushButton("Экспортировать в Word")
        self.export_button.setStyleSheet("background-color: #7ec25f; color: black;")
        self.export_button.clicked.connect(self.export_to_word)
        buttons_layout.addWidget(self.export_button, 3, 0, 1, 2)

        layout.addLayout(buttons_layout)

        self.result_label = QLabel("")
        layout.addWidget(self.result_label)

        self.chart_label = QLabel()
        layout.addWidget(self.chart_label)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.load_data("Sales", "Продажи")

    def load_data(self, table_name, display_name):
        cursor = self.connection.cursor()

        if table_name == "Sales":
            query = """
                SELECT 
                    s.SaleID AS 'ID продажи',
                    CONCAT(cust.LastName, ' ', cust.FirstName, ' ', cust.MiddleName) AS 'Клиент',
                    CONCAT(mgr.LastName, ' ', mgr.FirstName, ' ', mgr.MiddleName) AS 'Менеджер',
                    CONCAT(cars.Brand, ' ', cars.Model) AS 'Автомобиль',
                    s.LicensePlate AS 'Госномер',
                    s.SaleDate AS 'Дата продажи',
                    s.SalePrice AS 'Сумма продажи'
                FROM Sales s
                JOIN Customers cust ON s.CustomerID = cust.CustomerID
                JOIN Managers mgr ON s.ManagerID = mgr.ManagerID
                JOIN Cars cars ON s.CarID = cars.CarID
            """
            cursor.execute(query)
        else:
            cursor.execute(f"SELECT * FROM {table_name}")

        rows = cursor.fetchall()
        column_names = [desc[0] for desc in cursor.description]

        self.table.setRowCount(len(rows))
        self.table.setColumnCount(len(column_names))
        self.table.setHorizontalHeaderLabels(column_names)

        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                self.table.setItem(i, j, QTableWidgetItem(str(value)))

        self.result_label.setText(f"Загружены данные из таблицы: {display_name}")

    def export_to_word(self):
        try:
            row_count = self.table.rowCount()
            column_count = self.table.columnCount()

            doc = Document()
            doc.add_heading('Отчет: Данные таблицы', level=1)

            table = doc.add_table(rows=row_count + 1, cols=column_count)
            table.style = 'Table Grid'

            for col in range(column_count):
                cell = table.cell(0, col)
                cell.text = self.table.horizontalHeaderItem(col).text()

            for row in range(row_count):
                for col in range(column_count):
                    cell = table.cell(row + 1, col)
                    cell.text = self.table.item(row, col).text()

            file_path = os.path.join(os.getcwd(), "отчет_автосалон.docx")
            doc.save(file_path)
            QMessageBox.information(self, "Успех", f"Данные успешно экспортированы в файл: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте данных: {e}")

    def add_customer(self):
        self.add_record_form("Добавить клиента", ["Фамилия", "Имя", "Отчество"], self.insert_customer)

    def add_manager(self):
        self.add_record_form("Добавить менеджера", ["Фамилия", "Имя", "Отчество"], self.insert_manager)

    def add_car(self):
        self.add_record_form("Добавить автомобиль", ["Марка", "Модель", "Год выпуска", "Цена"], self.insert_car)

    def add_sale(self):
        cursor = self.connection.cursor()
        cursor.execute("SELECT CustomerID, CONCAT(LastName, ' ', FirstName) FROM Customers")
        customers = cursor.fetchall()

        cursor.execute("SELECT ManagerID, CONCAT(LastName, ' ', FirstName) FROM Managers")
        managers = cursor.fetchall()

        cursor.execute("SELECT CarID, CONCAT(Brand, ' ', Model) FROM Cars")
        cars = cursor.fetchall()

        if not customers or not managers or not cars:
            QMessageBox.warning(self, "Ошибка",
                                "Необходимо добавить клиентов, менеджеров и автомобили перед добавлением продаж.")
            return

        form_layout = QFormLayout()
        self.form_widgets = {}

        self.form_widgets['Клиент'] = QComboBox()
        self.form_widgets['Клиент'].addItems([f"{c[0]}: {c[1]}" for c in customers])
        form_layout.addRow("Клиент:", self.form_widgets['Клиент'])

        self.form_widgets['Менеджер'] = QComboBox()
        self.form_widgets['Менеджер'].addItems([f"{m[0]}: {m[1]}" for m in managers])
        form_layout.addRow("Менеджер:", self.form_widgets['Менеджер'])

        self.form_widgets['Автомобиль'] = QComboBox()
        self.form_widgets['Автомобиль'].addItems([f"{c[0]}: {c[1]}" for c in cars])
        form_layout.addRow("Автомобиль:", self.form_widgets['Автомобиль'])

        self.form_widgets['Госномер'] = QLineEdit()
        form_layout.addRow("Госномер:", self.form_widgets['Госномер'])

        self.form_widgets['Дата продажи'] = QLineEdit()
        form_layout.addRow("Дата продажи (YYYY-MM-DD):", self.form_widgets['Дата продажи'])

        self.form_widgets['Сумма продажи'] = QLineEdit()
        form_layout.addRow("Сумма продажи:", self.form_widgets['Сумма продажи'])

        self.show_form("Добавить продажу", form_layout, self.insert_sale)

    def add_record_form(self, title, fields, submit_function):
        form_layout = QFormLayout()
        self.form_widgets = {}

        for field in fields:
            self.form_widgets[field] = QLineEdit()
            form_layout.addRow(f"{field}:", self.form_widgets[field])

        self.show_form(title, form_layout, submit_function)

    def show_form(self, title, form_layout, submit_function):
        form_window = QWidget()
        form_window.setWindowTitle(title)
        form_window.setGeometry(150, 150, 400, 300)

        submit_button = QPushButton("Сохранить")
        submit_button.setStyleSheet("background-color: #7ec25f; color: black;")
        submit_button.clicked.connect(lambda: self.submit_form(form_window, submit_function))

        form_layout.addWidget(submit_button)

        form_window.setLayout(form_layout)
        form_window.show()

    def submit_form(self, form_window, submit_function):
        try:
            submit_function()
            QMessageBox.information(self, "Успех", "Данные успешно добавлены!")
            form_window.close()
            self.load_data("Sales", "Продажи")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка добавления данных: {e}")

    def insert_customer(self):
        cursor = self.connection.cursor()
        cursor.execute(
            "INSERT INTO Customers (LastName, FirstName, MiddleName) VALUES (%s, %s, %s)",
            (
                self.form_widgets["Фамилия"].text(),
                self.form_widgets["Имя"].text(),
                self.form_widgets["Отчество"].text(),
            )
        )
        self.connection.commit()

    def insert_manager(self):
        cursor = self.connection.cursor()
        cursor.execute(
            "INSERT INTO Managers (LastName, FirstName, MiddleName) VALUES (%s, %s, %s)",
            (
                self.form_widgets["Фамилия"].text(),
                self.form_widgets["Имя"].text(),
                self.form_widgets["Отчество"].text(),
            )
        )
        self.connection.commit()

    def insert_car(self):
        cursor = self.connection.cursor()
        cursor.execute(
            "INSERT INTO Cars (Brand, Model, Year, Price) VALUES (%s, %s, %s, %s)",
            (
                self.form_widgets["Марка"].text(),
                self.form_widgets["Модель"].text(),
                int(self.form_widgets["Год выпуска"].text()),
                float(self.form_widgets["Цена"].text()),
            )
        )
        self.connection.commit()

    def insert_sale(self):
        customer_id = int(self.form_widgets["Клиент"].currentText().split(":")[0])
        manager_id = int(self.form_widgets["Менеджер"].currentText().split(":")[0])
        car_id = int(self.form_widgets["Автомобиль"].currentText().split(":")[0])

        cursor = self.connection.cursor()
        cursor.execute(
            "INSERT INTO Sales (CustomerID, ManagerID, CarID, LicensePlate, SaleDate, SalePrice) VALUES (%s, %s, %s, %s, %s, %s)",
            (
                customer_id,
                manager_id,
                car_id,
                self.form_widgets["Госномер"].text(),
                self.form_widgets["Дата продажи"].text(),
                float(self.form_widgets["Сумма продажи"].text()),
            )
        )
        self.connection.commit()

    def calculate_average(self):
        cursor = self.connection.cursor()
        cursor.execute("SELECT AVG(SalePrice) FROM Sales")
        avg_price = cursor.fetchone()[0]
        self.result_label.setText(f"Средняя сумма сделки: {avg_price:.2f}")

    def calculate_share(self):
        cursor = self.connection.cursor()
        cursor.execute(
            "SELECT Brand, COUNT(*) AS Count FROM Sales s JOIN Cars c ON s.CarID = c.CarID GROUP BY Brand"
        )
        rows = cursor.fetchall()

        labels = [row[0] for row in rows]
        sizes = [row[1] for row in rows]

        plt.figure(figsize=(4, 3))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        plt.axis('equal')

        plt.gcf().set_facecolor('#e1f5d7')

        chart_file = os.path.join(os.getcwd(), "chart.png")
        plt.savefig(chart_file)

        self.chart_label.setPixmap(QPixmap(chart_file))
        self.result_label.setText("Доля продаж по маркам отображена в диаграмме.")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = AutoSalonApp()
    window.show()
    sys.exit(app.exec())
